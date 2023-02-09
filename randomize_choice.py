from docx import Document
import random
import sys


FIXED_CHOICE = [
    "all of the above",
    "none of the above",
    "neither statement",
    "both statement"
]


def randomize_choice(path, version_number=1, num_choices=4):
    doc = Document(path)
    new_doc = Document(path)

    paragraphs = doc.paragraphs
    lines_new_question = [-1] + [i for i, p in enumerate(paragraphs) if p.text.strip() == '']

    row = 0
    i = 1
    answers = []
    while i < len(lines_new_question):
        if lines_new_question[i] - lines_new_question[i-1] < num_choices + 2:
            i += 1
            continue
        if row < lines_new_question[i] - num_choices:
            row += 1
        else:
            fix = False
            for choice in FIXED_CHOICE:
                if choice in "".join([p.text for p in paragraphs[row:row + num_choices]]).lower():
                    fix = True
                    break
            if fix:
                rd_paras = paragraphs[row:row + num_choices]
            else:
                rd_paras = random.sample(paragraphs[row:row + num_choices], num_choices)
            colors = [p.runs[0].font.color.rgb for p in rd_paras]
            if colors[0] == colors[1]:
                for c in range(num_choices):
                    if colors[c] != colors[0]:
                        answers.append(c)
                        break
            else:
                if colors[0] == colors[2]:
                    answers.append(1)
                else:
                    answers.append(0)

            for j in range(num_choices):
                new_doc.paragraphs[row+j].text = rd_paras[j].text
                new_doc.paragraphs[row+j].runs[0].bold = rd_paras[j].runs[0].bold
                new_doc.paragraphs[row+j].runs[0].italic = rd_paras[j].runs[0].italic
                new_doc.paragraphs[row+j].runs[0].underline = rd_paras[j].runs[0].underline
                new_doc.paragraphs[row+j].runs[0].font.color.rgb = rd_paras[j].runs[0].font.color.rgb
                new_doc.paragraphs[row+j].runs[0].font.name = rd_paras[j].runs[0].font.name
                new_doc.paragraphs[row+j].runs[0].style.name = rd_paras[j].runs[0].style.name
            row += num_choices
            i += 1
    for j, p in enumerate(paragraphs[row+1:]):
        if p.runs[0].font.color.rgb:
            answers.append(num_choices + j - len(paragraphs[row + 1:]))

    if path[:-7].endswith("Version") and path[-6].isdigit():
        new_path = path[:-6] + str(int(path[-6]) + 1) + path[-5:]
    else:
        new_path = path[:-5] + f" - Version {version_number}" + path[-5:]

    with open(new_path[:-5] + " Answer key.txt", 'w') as f:
        f.write(''.join([{0: 'A', 1: 'B', 2: 'C', 3: 'D', 4: 'E'}[a] + ('\n' if i % 5 == 4 else '') for i, a in enumerate(answers)]))
    new_doc.save(new_path)


if __name__ == '__main__':
    args = sys.argv[1:]
    path = args[0]
    versions = 1
    num_choices = 4
    if len(args) > 1 and args[1].isdigit():
        versions = int(args[1])
    if len(args) > 2 and args[2].isdigit():
        num_choices = int(args[2])

    if path.endswith(".docx"):
        for i in range(versions):
            randomize_choice(path, i + 1, num_choices)
